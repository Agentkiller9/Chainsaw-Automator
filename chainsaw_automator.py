import os
import json
import csv
import subprocess
import sys
import argparse
from datetime import datetime
from pathlib import Path
import time
from collections import defaultdict
from typing import Dict, Set, Union, DefaultDict, Any, Optional, Callable, TypeVar

# Define a generic type for Document
T = TypeVar('T')
DocumentType = Callable[..., Any]

# Suppress import-related type checking warnings
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define type aliases for better readability
UserStats = Dict[str, Union[int, Set[str]]]
EventData = Dict[str, Any]

# Type definitions for python-docx
class DocxDocument:
    def add_heading(self, text: str, level: int) -> 'DocxHeading': ...
    def add_paragraph(self, text: str = '') -> 'DocxParagraph': ...
    def add_table(self, rows: int, cols: int) -> 'DocxTable': ...
    def save(self, path: str) -> None: ...

class DocxHeading:
    alignment: int

class DocxParagraph:
    def add_run(self, text: str) -> 'DocxRun': ...

class DocxTable:
    style: str
    rows: list['DocxTableRow']
    def add_row(self) -> 'DocxTableRow': ...

class DocxTableRow:
    cells: list['DocxTableCell']

class DocxTableCell:
    text: str

class DocxRun:
    pass

class WdAlignParagraph:
    CENTER: int = 1

# Global variables for docx module with proper type hints
DOCX_AVAILABLE = False
Document: Optional[DocumentType] = None
WD_ALIGN_PARAGRAPH: Any = None  # Using Any since the actual type is complex

# Attempt to import python-docx, with proper error handling
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not installed. Word report generation will be disabled.")

class ChainsawAnalyzer:
    def __init__(self):
        self.chainsaw_path = None
        self.evtx_folder = None
        self.output_folder = None
        self.server_name = None
        self.csv_format = 'expert'  # default
        # Predefined event mappings for common security events
        self.event_mappings = {
            "logon_events": {
                "4624": "Successful Logon",
                "4625": "Failed Logon",
                "4634": "Account Logoff",
                "4647": "User Logoff",
                "4778": "Session Reconnected",
                "4779": "Session Disconnected"
            },
            "account_management": {
                "4720": "Account Created",
                "4722": "Account Enabled",
                "4723": "Account Disabled",
                "4724": "Account Password Changed",
                "4725": "Account Password Reset",
                "4728": "Member Added to Security Group",
                "4729": "Member Removed from Security Group"
            },
            "process_events": {
                "4688": "Process Created",
                "4689": "Process Terminated"
            },
            "network_events": {
                "4624": "Network Logon",
                "4625": "Network Logon Failure"
            },
            "file_events": {
                "4663": "File Accessed",
                "4660": "Object Deleted",
                "4661": "Object Deleted"
            }
        }
        # LogonType mapping
        self.logon_type_mapping = {
            "2": "Interactive",
            "3": "Network",
            "4": "Batch",
            "5": "Service",
            "7": "Unlock",
            "8": "NetworkCleartext",
            "9": "NewCredentials",
            "10": "RemoteInteractive",
            "11": "CachedInteractive"
        }

    def setup_paths(self):
        """Interactive setup of paths and parameters"""
        print("🔍 Enhanced Chainsaw Log Analysis Tool")
        print("=" * 60)
        # Get chainsaw path
        self.chainsaw_path = input("Enter the path to chainsaw.exe: ").strip('"')
        if not self.chainsaw_path:
            self.chainsaw_path = r"C:\Users\mujta\Downloads\Tools\chainsaw\chainsaw.exe"
            print(f"Using default path: {self.chainsaw_path}")
        # Verify chainsaw exists
        if not self.chainsaw_path or not os.path.exists(self.chainsaw_path):
            print(f"❌ Chainsaw not found at: {self.chainsaw_path}")
            return False
        print(f"✅ Chainsaw found at: {self.chainsaw_path}")
        # Get EVTX folder
        self.evtx_folder = input("Enter the path to the folder containing EVTX files: ").strip('"')
        if not self.evtx_folder or not os.path.exists(self.evtx_folder):
            print(f"❌ EVTX folder not found: {self.evtx_folder}")
            return False
        print(f"✅ EVTX folder found: {self.evtx_folder}")
        # Get server name
        self.server_name = input("Enter the server name (for output file naming): ").strip()
        if not self.server_name:
            self.server_name = "default_server"
        # Get output folder
        output_folder = input("Enter the path for results folder (or press Enter for default): ").strip('"')
        if not output_folder:
            self.output_folder = os.path.join(os.path.dirname(self.chainsaw_path), "results", self.server_name)
        else:
            self.output_folder = output_folder
        if not self.output_folder:
            print("❌ Output folder could not be determined.")
            return False
        # Ask for CSV format
        print("\nCSV Output Format Options:")
        print("1. Simple (minimal fields, like AnalyzeChainsaw_CustomOutput.py)")
        print("2. Expert (comprehensive, all fields)")
        format_choice = input("Choose CSV format (1 for Simple, 2 for Expert) [2]: ").strip()
        if format_choice == "1":
            self.csv_format = 'simple'
        else:
            self.csv_format = 'expert'
        return True

    def show_event_categories(self):
        print("\n📋 Available Event Categories:")
        print("-" * 40)
        for category, events in self.event_mappings.items():
            print(f"\n🔹 {category.replace('_', ' ').title()}:")
            for event_id, description in events.items():
                print(f"   {event_id}: {description}")

    def get_custom_events(self):
        print("\n🎯 Custom Event Selection:")
        print("-" * 30)
        print("1. Use predefined categories")
        print("2. Enter custom event IDs")
        print("3. Use all common security events")
        choice = input("\nSelect option (1-3): ").strip()
        if choice == "1":
            return self.select_predefined_events()
        elif choice == "2":
            return self.get_manual_events()
        elif choice == "3":
            return self.get_all_security_events()
        else:
            print("Invalid choice, using default logon events")
            return {"4624": "Successful Logon", "4625": "Failed Logon"}

    def select_predefined_events(self):
        print("\n📋 Select Event Categories:")
        categories = list(self.event_mappings.keys())
        for i, category in enumerate(categories, 1):
            print(f"{i}. {category.replace('_', ' ').title()}")
        print(f"{len(categories) + 1}. All categories")
        try:
            choice = int(input(f"\nSelect category (1-{len(categories) + 1}): "))
            if choice == len(categories) + 1:
                return self.get_all_security_events()
            elif 1 <= choice <= len(categories):
                selected_category = categories[choice - 1]
                return self.event_mappings[selected_category]
            else:
                print("Invalid choice, using logon events")
                return self.event_mappings["logon_events"]
        except ValueError:
            print("Invalid input, using logon events")
            return self.event_mappings["logon_events"]

    def get_manual_events(self):
        events = {}
        print("\n📝 Enter Event IDs (press Enter when done):")
        while True:
            event_id = input("Event ID (or Enter to finish): ").strip()
            if not event_id:
                break
            description = input(f"Description for Event {event_id}: ").strip()
            if not description:
                description = f"Event {event_id}"
            events[event_id] = description
        return events if events else {"4624": "Successful Logon", "4625": "Failed Logon"}

    def get_all_security_events(self):
        all_events = {}
        for category in self.event_mappings.values():
            all_events.update(category)
        return all_events

    def run_chainsaw_search(self, event_id, description):
        try:
            if not self.output_folder:
                raise ValueError("Output folder is not set.")
            output_folder = self.output_folder or "results"
            os.makedirs(output_folder, exist_ok=True)
            server_name = self.server_name or "default_server"
            output_json = os.path.join(output_folder, f"{server_name}_{event_id}.json")
            if not self.chainsaw_path or not self.evtx_folder:
                raise ValueError("Chainsaw path or EVTX folder is not set.")
            command = [
                self.chainsaw_path,
                "search",
                "-t", f"Event.System.EventID: ={event_id}",
                self.evtx_folder,
                "--json"
            ]
            print(f"🔍 Searching for {description} (Event ID {event_id})...")
            with open(output_json, 'w', encoding='utf-8') as output_file:
                result = subprocess.run(command, stdout=output_file, stderr=subprocess.PIPE, 
                                      text=True, encoding='utf-8', errors='replace', check=True)
            if os.path.exists(output_json) and os.path.getsize(output_json) > 0:
                file_size = os.path.getsize(output_json)
                print(f"✅ {description}: {file_size} bytes")
                return output_json
            else:
                print(f"⚠  {description}: No results found")
                return None
        except subprocess.CalledProcessError as e:
            print(f"❌ Error searching for {description}: {e.stderr}")
            return None
        except Exception as e:
            print(f"❌ Unexpected error for {description}: {e}")
            return None

    def analyze_json_files(self, output_filename):
        output_folder = self.output_folder or "results"
        json_files = [f for f in os.listdir(output_folder) if f.endswith(".json")]
        if not json_files:
            print("❌ No JSON files found for analysis!")
            return
        print(f"\n📊 Analyzing {len(json_files)} JSON files...")
        all_events = []
        for file_name in json_files:
            file_path = os.path.join(output_folder, file_name)
            print(f"📄 Processing: {file_name}")
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                for entry in data:
                    if self.csv_format == 'simple':
                        event_data = self.extract_simple_event_data(entry, file_name)
                    else:
                        event_data = self.extract_event_data(entry, file_name)
                    if event_data:
                        all_events.append(event_data)
            except Exception as e:
                print(f"❌ Failed to process {file_name}: {e}")
                continue
        if not all_events:
            print("❌ No events found in JSON files!")
            return
        # Create CSV in selected format
        if self.csv_format == 'simple':
            self.create_simple_csv(all_events, output_filename)
        else:
            self.create_comprehensive_csv(all_events, output_filename)
        # Create summary report (always expert style)
        self.create_summary_report(all_events, output_filename)

    def extract_event_data(self, entry, source_file):
        try:
            system = entry.get("Event", {}).get("System", {})
            event_id = str(system.get("EventID", ""))
            time_created = system.get("TimeCreated_attributes", {}).get("SystemTime", "")
            event_data = entry.get("Event", {}).get("EventData", {})
            if isinstance(event_data, dict):
                data_dict = event_data
            elif isinstance(event_data, list):
                data_dict = {d.get("Name", ""): d.get("Value", "") for d in event_data}
            else:
                data_dict = {}
            extracted_data = {
                "EventID": event_id,
                "TimeCreated": time_created,
                "SourceFile": source_file,
                "TargetUserName": data_dict.get("TargetUserName", ""),
                "TargetDomainName": data_dict.get("TargetDomainName", ""),
                "TargetUserSid": data_dict.get("TargetUserSid", ""),
                "LogonType": data_dict.get("LogonType", ""),
                "LogonTypeName": self.logon_type_mapping.get(str(data_dict.get("LogonType", "")), "Unknown"),
                "IpAddress": data_dict.get("IpAddress", ""),
                "WorkstationName": data_dict.get("WorkstationName", ""),
                "Status": data_dict.get("Status", ""),
                "ProcessName": data_dict.get("ProcessName", ""),
                "NewUacValue": data_dict.get("NewUacValue", ""),
                "SubjectUserName": data_dict.get("SubjectUserName", ""),
                "SubjectDomainName": data_dict.get("SubjectDomainName", ""),
                "SubjectUserSid": data_dict.get("SubjectUserSid", ""),
                "ObjectName": data_dict.get("ObjectName", ""),
                "ObjectType": data_dict.get("ObjectType", ""),
                "AccessMask": data_dict.get("AccessMask", ""),
                "ObjectServer": data_dict.get("ObjectServer", ""),
                "HandleId": data_dict.get("HandleId", ""),
                "ProcessId": data_dict.get("ProcessId", ""),
                "Image": data_dict.get("Image", ""),
                "CommandLine": data_dict.get("CommandLine", ""),
                "CurrentDirectory": data_dict.get("CurrentDirectory", ""),
                "User": data_dict.get("User", ""),
                "LogonGuid": data_dict.get("LogonGuid", ""),
                "TransmittedServices": data_dict.get("TransmittedServices", ""),
                "LmPackageName": data_dict.get("LmPackageName", ""),
                "KeyLength": data_dict.get("KeyLength", ""),
                "ImpersonationLevel": data_dict.get("ImpersonationLevel", ""),
                "RestrictedAdminMode": data_dict.get("RestrictedAdminMode", ""),
                "TargetOutboundUserName": data_dict.get("TargetOutboundUserName", ""),
                "TargetOutboundDomainName": data_dict.get("TargetOutboundDomainName", ""),
                "VirtualAccount": data_dict.get("VirtualAccount", ""),
                "TargetLinkedLogonId": data_dict.get("TargetLinkedLogonId", ""),
                "ElevatedToken": data_dict.get("ElevatedToken", ""),
                "SourceNetworkAddress": data_dict.get("SourceNetworkAddress", ""),
                "SourcePort": data_dict.get("SourcePort", ""),
                "NewTargetUserName": data_dict.get("NewTargetUserName", ""),
                "NewTargetDomainName": data_dict.get("NewTargetDomainName", ""),
                "NewTargetUserSid": data_dict.get("NewTargetUserSid", ""),
                "SidHistory": data_dict.get("SidHistory", ""),
                "SamAccountName": data_dict.get("SamAccountName", ""),
                "DisplayName": data_dict.get("DisplayName", ""),
                "UserPrincipalName": data_dict.get("UserPrincipalName", ""),
                "HomeDirectory": data_dict.get("HomeDirectory", ""),
                "HomePath": data_dict.get("HomePath", ""),
                "ScriptPath": data_dict.get("ScriptPath", ""),
                "ProfilePath": data_dict.get("ProfilePath", ""),
                "UserWorkstations": data_dict.get("UserWorkstations", ""),
                "PasswordLastSet": data_dict.get("PasswordLastSet", ""),
                "AccountExpires": data_dict.get("AccountExpires", ""),
                "PrimaryGroupId": data_dict.get("PrimaryGroupId", ""),
                "AllowedToDelegateTo": data_dict.get("AllowedToDelegateTo", ""),
                "OldUacValue": data_dict.get("OldUacValue", ""),
                "UserAccountControl": data_dict.get("UserAccountControl", ""),
                "UserParameters": data_dict.get("UserParameters", ""),
                "LogonHours": data_dict.get("LogonHours", ""),
                "DnsHostName": data_dict.get("DnsHostName", ""),
                "ServicePrincipalNames": data_dict.get("ServicePrincipalNames", ""),
            }
            return extracted_data
        except Exception as e:
            print(f"❌ Error extracting data from event: {e}")
            return None

    def extract_simple_event_data(self, entry, source_file):
        try:
            system = entry.get("Event", {}).get("System", {})
            event_id = str(system.get("EventID", ""))
            time_created = system.get("TimeCreated_attributes", {}).get("SystemTime", "")
            event_data = entry.get("Event", {}).get("EventData", {})
            if isinstance(event_data, dict):
                data_dict = event_data
            elif isinstance(event_data, list):
                data_dict = {d.get("Name", ""): d.get("Value", "") for d in event_data}
            else:
                data_dict = {}
            target_user = data_dict.get("TargetUserName", "")
            logon_type = data_dict.get("LogonType", "")
            ip_address = data_dict.get("IpAddress", "")
            status = data_dict.get("Status", "")
            logon_type_name = self.logon_type_mapping.get(str(logon_type), "Unknown")
            return {
                "User": target_user,
                "EventID": event_id,
                "LogonType": logon_type,
                "LogonTypeName": logon_type_name,
                "Status": status,
                "IpAddress": ip_address,
                "TimeCreated": time_created,
                "SourceFile": source_file
            }
        except Exception as e:
            print(f"❌ Error extracting simple data from event: {e}")
            return None

    def create_comprehensive_csv(self, all_events, output_filename):
        output_folder = self.output_folder or "results"
        output_file = os.path.join(output_folder, output_filename)
        if not all_events:
            return
        fieldnames = set()
        for event in all_events:
            fieldnames.update(event.keys())
        fieldnames = sorted(list(fieldnames))
        with open(output_file, "w", encoding="utf-8", newline="") as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for event in all_events:
                writer.writerow(event)
        print(f"✅ Comprehensive CSV saved: {output_file}")
        print(f"📊 Total events: {len(all_events)}")
        print(f"📋 Fields extracted: {len(fieldnames)}")

    def create_simple_csv(self, all_events, output_filename):
        output_folder = self.output_folder or "results"
        output_file = os.path.join(output_folder, output_filename)
        if not all_events:
            return
        fieldnames = ["User", "EventID", "LogonType", "LogonTypeName", "Status", "IpAddress", "TimeCreated", "SourceFile"]
        with open(output_file, "w", encoding="utf-8", newline="") as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for event in all_events:
                writer.writerow({k: event.get(k, "") for k in fieldnames})
        print(f"✅ Simple CSV saved: {output_file}")
        print(f"📊 Total events: {len(all_events)}")

    def create_summary_report(self, all_events, output_filename):
        output_folder = self.output_folder or "results"
        report_file = os.path.join(output_folder, f"{output_filename.replace('.csv', '_summary.txt')}")
        with open(report_file, "w", encoding="utf-8") as f:
            f.write("🔍 CHAINSAW ANALYSIS SUMMARY REPORT\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Server: {self.server_name or 'default_server'}\n")
            f.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Total Events: {len(all_events)}\n\n")
            event_counts = {}
            for event in all_events:
                event_id = event.get("EventID", "Unknown")
                event_counts[event_id] = event_counts.get(event_id, 0) + 1
            f.write("📊 EVENT ID BREAKDOWN:\n")
            f.write("-" * 30 + "\n")
            for event_id, count in sorted(event_counts.items()):
                f.write(f"Event {event_id}: {count} events\n")
            logon_counts = {}
            for event in all_events:
                logon_type = event.get("LogonTypeName", "Unknown")
                logon_counts[logon_type] = logon_counts.get(logon_type, 0) + 1
            f.write("\n🔐 LOGON TYPE BREAKDOWN:\n")
            f.write("-" * 30 + "\n")
            for logon_type, count in sorted(logon_counts.items()):
                f.write(f"{logon_type}: {count} events\n")
            user_counts = {}
            for event in all_events:
                user = event.get("TargetUserName", event.get("User", "Unknown"))
                if user and user != "Unknown":
                    user_counts[user] = user_counts.get(user, 0) + 1
            f.write("\n👤 TOP USERS:\n")
            f.write("-" * 30 + "\n")
            top_users = sorted(user_counts.items(), key=lambda x: x[1], reverse=True)[:10]
            for user, count in top_users:
                f.write(f"{user}: {count} events\n")
        print(f"📄 Summary report saved: {report_file}")

    def run_analysis(self):
        if not self.setup_paths():
            return
        self.show_event_categories()
        events_to_search = self.get_custom_events()
        if not events_to_search:
            print("❌ No events selected for analysis!")
            return
        print(f"\n🎯 Searching for {len(events_to_search)} event types...")
        output_filename = input("Enter the desired name for the final CSV summary (without .csv extension): ").strip()
        if not output_filename.endswith('.csv'):
            output_filename += '.csv'
        print("\n🚀 Starting analysis...")
        print("=" * 60)
        successful_searches = []
        for event_id, description in events_to_search.items():
            result = self.run_chainsaw_search(event_id, description)
            if result:
                successful_searches.append(result)
        if not successful_searches:
            print("\n❌ No successful searches completed!")
            return
        print(f"\n📋 Analyzing {len(successful_searches)} result files...")
        self.analyze_json_files(output_filename)
        print("\n🎉 Analysis complete!")
        output_folder = self.output_folder or "results"
        print(f"📁 Results folder: {output_folder}")
        print(f"📄 JSON files created: {len(successful_searches)}")
        print(f"📊 Summary CSV: {os.path.join(output_folder, output_filename)}")

def main():
    analyzer = ChainsawAnalyzer()
    analyzer.run_analysis()

if __name__ == "__main__":
    main()
